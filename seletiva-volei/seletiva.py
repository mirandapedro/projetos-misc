import csv
import docx

with open("docs\\dados-seletiva.csv", "r", encoding="utf-8-sig") as f:
    dados_seletiva = [aluno for aluno in csv.DictReader(f, delimiter=";")]
    
for aluno in dados_seletiva:
    if aluno["Ano Nascimento"] in ["2005", "2006"]:
        aluno["ordem"] = 2
    elif aluno["Ano Nascimento"] in ["2007", "2008", "2009", "2010"]:
        aluno["ordem"] = 1
    else:
        print(aluno)

dados_seletiva = sorted(dados_seletiva, key=lambda x: x["Nome"])
dados_seletiva = sorted(dados_seletiva, key=lambda x: x["ordem"])

with open("docs\\modelo-criterios.docx", "rb") as f:
    documento = docx.Document(f)

table = documento.add_table(len(dados_seletiva), 9)
linha_atual = 0
for aluno in dados_seletiva:
    cell = table.cell(linha_atual, 0)
    cell.paragraphs[0].add_run()
    cell.paragraphs[0].runs[0].add_picture("fotos-alunos\\{}.jpeg".format(aluno["ID"]),
                                         width = docx.shared.Cm(1.5))
    cell = table.cell(linha_atual, 1)
    cell.add_paragraph("{}\n{}\n{}\n{}".format(aluno["Nome"],
                                               aluno["Ano Nascimento"],
                                               aluno["Posição"],
                                               aluno["Altura"]))
    
    linha_atual += 1

documento.save("docs\\Critérios.docx")