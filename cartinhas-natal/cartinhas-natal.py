import docx
import csv
import os
import docx2pdf
import fitz
import datetime

def remove_espaco(str_espaco):
    if len(str_espaco) == 0: return str_espaco
    nova_str = str_espaco
    while nova_str[0] == " ":
        if nova_str == " ": return ""
        nova_str = nova_str[1:]
    while nova_str[-1] == " ":
        nova_str = nova_str[:-1]
    return nova_str
'''
def calcula_idade(nascimento):
    idade = datetime.date(int(nascimento.split("/")[2]), int(nascimento.split("/")[1]),
                          int(nascimento.split("/")[0]))
    print(datetime.datetime.today())
    idade = datetime.datetime.today() - idade
    print(idade)
'''    
def trata_dados(dados_criancas):
    novos_dados = dados_criancas
    novas_criancas = []
    for crianca in novos_dados:
        if crianca == {}: continue
        for key in ["criança", "nascimento", "calçado", "camisa", "calça", "sexo"]:
            crianca[key] = remove_espaco(str(crianca[key]))
        for key in ["calçado", "camisa", "calça"]:
            crianca[key] = crianca[key].replace("a", " anos").replace("m", " meses")
        #crianca["idade"] = calcula_idade(crianca["nascimento"])
        novas_criancas.append(crianca)
    return novas_criancas

def verifica_completo(crianca):
    if crianca == {}: return False
    if not os.path.isfile("fotos\\_{}.jpg".format(crianca["criança"])):
        print("Não encontrei a foto dx {}!".format(crianca["criança"]))
        return False
    keys_importantes = ["criança", "nascimento", "calçado", "camisa", "calça", "sexo"]
    if any(len(crianca[x]) == 0 for x in keys_importantes):
        print("{} está com dados incompletos!".format(crianca["criança"]))
        return False
    return True

def adiciona_intro(document, sexo):
    novo_document = document
    if sexo == "M": corpo = "Este é o seu afilhado no Natal 2023. Presenteie-o"
    elif sexo == "F": corpo = "Esta é a sua afilhada no Natal 2023. Presenteie-a"
    run = novo_document.paragraphs[1].runs[0]
    run.text = run.text.format(corpo)
    return novo_document

def adiciona_foto(table, document, caminho):
    nova_table = table
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run()
    run.add_picture(caminho, height=docx.shared.Cm(8))
    cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    return nova_table

def trata_nome(str_nome):
    return str_nome \
            .replace(" ", "-") \
            .replace("á", "a") \
            .replace("ã", "a") \
            .replace("à", "a") \
            .replace("é", "e") \
            .replace("ê", "e") \
            .replace("í", "i") \
            .replace("ó", "o") \
            .replace("ô", "o") \
            .replace("ú", "u") \
            .replace("ü", "u") \
            .replace("ç", "c")

def adiciona_dados(table, document, crianca):
    nova_table = table
    nova_table.cell(0, 1).paragraphs[0].text = crianca["criança"]
    nova_table.cell(0, 1).paragraphs[0].style = document.styles["Normal"]
    nova_table.cell(0, 1).paragraphs[0].runs[0].font.size = docx.shared.Pt(24)
    nova_table.cell(0, 1).paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    nova_table.cell(0, 1).add_paragraph()
    nova_table.cell(1, 1).paragraphs[0].text = crianca["nascimento"]
    nova_table.cell(1, 1).paragraphs[0].style = document.styles["Normal"]
    run = nova_table.cell(2, 1).paragraphs[0].add_run(crianca["camisa"])
    run.bold = True
    run = nova_table.cell(3, 1).paragraphs[0].add_run(crianca["calça"])
    run.bold = True
    run = nova_table.cell(4, 1).paragraphs[0].add_run(crianca["calçado"])
    run.bold = True
    run = nova_table.cell(5, 1).paragraphs[0].add_run(crianca["sexo"])
    for row in range(1, 6):
        for run in nova_table.cell(row, 1).paragraphs[0].runs:    
            run.font.size = docx.shared.Pt(18)
    return nova_table
    

with open("Planilha Natal 2023 - Sheet1.csv", "r", encoding="utf-8") as f:
    dados_criancas = [x for x in csv.DictReader(f)]
dados_criancas = trata_dados(dados_criancas)

#TODO conferir fotos na pasta que ainda estão sem donos

with open ("docs\\modelo.docx", "rb") as f:
    for crianca in dados_criancas:
        if not verifica_completo(crianca):
            continue
        document = docx.Document(f)
        document = adiciona_intro(document, crianca["sexo"])
        table = document.tables[-1]
        table = adiciona_foto(table, document, "fotos\\_{}.jpg"\
                              .format(crianca["criança"]))
        table = adiciona_dados(table, document, crianca)
        document.save("cartinhas\\docx\\{}.docx"\
                      .format(trata_nome(crianca["criança"])))

for file in os.listdir("cartinhas\\docx"):
    if "~" in file: continue
    docx2pdf.convert(os.path.abspath("cartinhas/docx/"+file), 
                     os.path.abspath("cartinhas/pdf/"+file[:-4]+"pdf"))
    doc = fitz.open("cartinhas\\pdf\\"+file[:-4]+"pdf")
    page = doc.load_page(0)  # number of page
    pix = page.get_pixmap(matrix=fitz.Matrix(300/72,300/72))
    pix.save("cartinhas\\img\\" + file[:-4]+"png")


document = docx.Document()
table = document.add_table(1, 2)
linha_atual = 0
dados_criancas = sorted(dados_criancas, key = lambda x: x["criança"])
for crianca in dados_criancas:
    if not verifica_completo(crianca):
        continue
    cell = table.cell(linha_atual, 0)
    cell.paragraphs[0].add_run()
    cell.paragraphs[0].runs[0].add_picture("fotos\\_{}.jpg".format(crianca["criança"]),
                                           height = docx.shared.Cm(4))
    cell = table.cell(linha_atual, 1)
    cell.text = crianca["criança"]
    table.add_row()
    linha_atual += 1
document.save("etiquetas.docx")
